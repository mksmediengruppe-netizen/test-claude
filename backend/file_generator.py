"""
File Generator — модуль генерации файлов для Super Agent v5.0.

Поддерживаемые форматы:
- .txt — простой текст
- .md — Markdown
- .html — HTML страницы
- .json — JSON данные
- .csv — CSV таблицы
- .docx — Word документы (python-docx)
- .pdf — PDF документы (fpdf2)
- .xlsx — Excel таблицы (openpyxl)
- .py / .js / .css / .sql — код

Каждый файл сохраняется в GENERATED_DIR и получает уникальный ID для скачивания.
"""

import os
import json
import uuid
import time
import hashlib
import logging
from datetime import datetime, timezone

logger = logging.getLogger("file_generator")

# Directory for generated files
# Auto-detect project dir for generated files
_PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
GENERATED_DIR = os.environ.get("GENERATED_DIR", os.path.join(_PROJECT_DIR, "generated"))
os.makedirs(GENERATED_DIR, exist_ok=True)

# Registry of generated files (in-memory, persisted to JSON)
_registry_path = os.path.join(GENERATED_DIR, "_registry.json")
_registry = {}


def _load_registry():
    global _registry
    try:
        if os.path.exists(_registry_path):
            with open(_registry_path, "r") as f:
                _registry = json.load(f)
    except Exception:
        _registry = {}


def _save_registry():
    try:
        with open(_registry_path, "w") as f:
            json.dump(_registry, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Failed to save registry: {e}")


def _register_file(file_id, filename, filepath, format_type, size, chat_id=None, user_id=None):
    """Register a generated file in the registry."""
    _load_registry()
    _registry[file_id] = {
        "id": file_id,
        "filename": filename,
        "filepath": filepath,
        "format": format_type,
        "size": size,
        "chat_id": chat_id,
        "user_id": user_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "downloads": 0
    }
    _save_registry()
    return _registry[file_id]


def get_file_info(file_id):
    """Get file info by ID."""
    _load_registry()
    return _registry.get(file_id)


def get_file_path(file_id):
    """Get file path by ID for serving."""
    _load_registry()
    info = _registry.get(file_id)
    if info and os.path.exists(info["filepath"]):
        # Increment download counter
        info["downloads"] += 1
        _save_registry()
        return info["filepath"], info["filename"]
    return None, None


def list_files(chat_id=None, user_id=None, limit=50):
    """List generated files, optionally filtered."""
    _load_registry()
    files = list(_registry.values())
    if chat_id:
        files = [f for f in files if f.get("chat_id") == chat_id]
    if user_id:
        files = [f for f in files if f.get("user_id") == user_id]
    files.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return files[:limit]


# ══════════════════════════════════════════════════════════════
# FILE GENERATION FUNCTIONS
# ══════════════════════════════════════════════════════════════

def generate_text_file(content, filename="document.txt", chat_id=None, user_id=None):
    """Generate a plain text file."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "txt", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_markdown_file(content, filename="document.md", chat_id=None, user_id=None):
    """Generate a Markdown file."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "md", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_html_file(content, filename="page.html", chat_id=None, user_id=None):
    """Generate an HTML file."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "html", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download",
            "preview_url": f"/api/files/{file_id}/preview"}


def generate_json_file(data, filename="data.json", chat_id=None, user_id=None):
    """Generate a JSON file."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    content = json.dumps(data, indent=2, ensure_ascii=False) if isinstance(data, (dict, list)) else str(data)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "json", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_csv_file(content, filename="data.csv", chat_id=None, user_id=None):
    """Generate a CSV file."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "csv", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_code_file(content, filename="script.py", chat_id=None, user_id=None):
    """Generate a code file (.py, .js, .css, .sql, etc.)."""
    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")
    ext = os.path.splitext(filename)[1].lstrip(".")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, ext, size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_docx_file(content, filename="document.docx", title=None, chat_id=None, user_id=None):
    """Generate a Word .docx file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Install with: pip install python-docx"}

    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Add title if provided
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse content — support basic markdown-like formatting
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue

        # Headers
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        # Bullet lists
        elif stripped.startswith('- ') or stripped.startswith('• '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        # Numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
            doc.add_paragraph(stripped[2:].strip(), style='List Number')
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            # Handle bold **text**
            parts = stripped.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True

    doc.save(filepath)
    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "docx", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_pdf_file(content, filename="document.pdf", title=None, chat_id=None, user_id=None):
    """Generate a PDF file using fpdf2 with Unicode support."""
    try:
        from fpdf import FPDF
    except ImportError:
        return {"success": False, "error": "fpdf2 not installed. Install with: pip install fpdf2"}

    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Find a Unicode font
    font_name = "Helvetica"
    font_paths = [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVu"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "Liberation"),
        ("/usr/share/fonts/truetype/freefont/FreeSans.ttf",
         "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf", "FreeSans"),
    ]
    for regular, bold, name in font_paths:
        if os.path.exists(regular):
            try:
                pdf.add_font(name, "", regular, uni=True)
                if os.path.exists(bold):
                    pdf.add_font(name, "B", bold, uni=True)
                font_name = name
                break
            except Exception:
                continue

    pdf.set_font(font_name, size=11)

    # Title
    if title:
        pdf.set_font(font_name, "B", 18)
        pdf.cell(0, 12, title, ln=True, align="C")
        pdf.ln(8)
        pdf.set_font(font_name, size=11)

    # Content — replace problematic chars for non-unicode fonts
    def safe_text(t):
        if font_name == "Helvetica":
            return t.replace('\u2022', '-').replace('\u2014', '--').replace('\u2013', '-').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
        return t

    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            pdf.ln(4)
            continue

        if stripped.startswith('# '):
            pdf.set_font(font_name, "B", 16)
            pdf.cell(0, 10, safe_text(stripped[2:]), ln=True)
            pdf.set_font(font_name, size=11)
        elif stripped.startswith('## '):
            pdf.set_font(font_name, "B", 14)
            pdf.cell(0, 9, safe_text(stripped[3:]), ln=True)
            pdf.set_font(font_name, size=11)
        elif stripped.startswith('### '):
            pdf.set_font(font_name, "B", 12)
            pdf.cell(0, 8, safe_text(stripped[4:]), ln=True)
            pdf.set_font(font_name, size=11)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            bullet = '- ' if font_name == 'Helvetica' else '\u2022 '
            pdf.cell(10)
            pdf.multi_cell(0, 6, safe_text(f"{bullet}{stripped[2:]}"))
        else:
            clean = stripped.replace('**', '')
            pdf.multi_cell(0, 6, safe_text(clean))

    pdf.output(filepath)
    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "pdf", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


def generate_xlsx_file(data, filename="spreadsheet.xlsx", sheet_name="Sheet1",
                       headers=None, chat_id=None, user_id=None):
    """
    Generate an Excel .xlsx file.
    data: list of lists (rows) or list of dicts
    headers: optional list of column headers
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        return {"success": False, "error": "openpyxl not installed. Install with: pip install openpyxl"}

    file_id = str(uuid.uuid4())[:12]
    filepath = os.path.join(GENERATED_DIR, f"{file_id}_{filename}")

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # Style for headers
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="6366F1", end_color="6366F1", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Handle different data formats
    if isinstance(data, str):
        # Parse CSV-like string
        rows = []
        for line in data.strip().split('\n'):
            if ',' in line:
                rows.append([cell.strip().strip('"') for cell in line.split(',')])
            elif '\t' in line:
                rows.append([cell.strip() for cell in line.split('\t')])
            else:
                rows.append([line.strip()])

        if rows:
            # First row as headers if no explicit headers
            if not headers and len(rows) > 1:
                headers = rows[0]
                rows = rows[1:]

            if headers:
                for col, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border

            for r, row in enumerate(rows, 2 if headers else 1):
                for c, val in enumerate(row, 1):
                    cell = ws.cell(row=r, column=c, value=val)
                    cell.border = thin_border

    elif isinstance(data, list):
        if data and isinstance(data[0], dict):
            # List of dicts
            if not headers:
                headers = list(data[0].keys())

            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border

            for r, item in enumerate(data, 2):
                for c, h in enumerate(headers, 1):
                    cell = ws.cell(row=r, column=c, value=item.get(h, ""))
                    cell.border = thin_border

        elif data and isinstance(data[0], (list, tuple)):
            # List of lists
            if headers:
                for col, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border

            start_row = 2 if headers else 1
            for r, row in enumerate(data, start_row):
                for c, val in enumerate(row, 1):
                    cell = ws.cell(row=r, column=c, value=val)
                    cell.border = thin_border

    # Auto-fit column widths (approximate)
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max_length + 4, 50)

    wb.save(filepath)
    size = os.path.getsize(filepath)
    info = _register_file(file_id, filename, filepath, "xlsx", size, chat_id, user_id)
    return {"success": True, "file_id": file_id, "filename": filename, "size": size,
            "download_url": f"/api/files/{file_id}/download"}


# ══════════════════════════════════════════════════════════════
# UNIVERSAL GENERATE FUNCTION
# ══════════════════════════════════════════════════════════════

def generate_file(content, filename, format_type=None, title=None,
                  headers=None, sheet_name="Sheet1",
                  chat_id=None, user_id=None):
    """
    Universal file generation function.
    Automatically detects format from filename extension.
    """
    if not format_type:
        _, ext = os.path.splitext(filename.lower())
        format_type = ext.lstrip(".")

    format_type = format_type.lower()

    generators = {
        "txt": lambda: generate_text_file(content, filename, chat_id, user_id),
        "md": lambda: generate_markdown_file(content, filename, chat_id, user_id),
        "markdown": lambda: generate_markdown_file(content, filename, chat_id, user_id),
        "html": lambda: generate_html_file(content, filename, chat_id, user_id),
        "htm": lambda: generate_html_file(content, filename, chat_id, user_id),
        "json": lambda: generate_json_file(content, filename, chat_id, user_id),
        "csv": lambda: generate_csv_file(content, filename, chat_id, user_id),
        "docx": lambda: generate_docx_file(content, filename, title, chat_id, user_id),
        "doc": lambda: generate_docx_file(content, filename.replace('.doc', '.docx'), title, chat_id, user_id),
        "pdf": lambda: generate_pdf_file(content, filename, title, chat_id, user_id),
        "xlsx": lambda: generate_xlsx_file(content, filename, sheet_name, headers, chat_id, user_id),
        "xls": lambda: generate_xlsx_file(content, filename.replace('.xls', '.xlsx'), sheet_name, headers, chat_id, user_id),
    }

    # Code files
    code_exts = {"py", "js", "ts", "css", "scss", "sql", "sh", "bash", "yaml", "yml",
                 "toml", "ini", "conf", "nginx", "xml", "jsx", "tsx", "vue", "svelte",
                 "rb", "go", "rs", "java", "kt", "swift", "php", "c", "cpp", "h"}

    if format_type in code_exts:
        return generate_code_file(content, filename, chat_id, user_id)

    gen = generators.get(format_type)
    if gen:
        return gen()

    # Fallback: save as text
    return generate_text_file(content, filename, chat_id, user_id)


def cleanup_old_files(max_age_hours=72):
    """Remove generated files older than max_age_hours."""
    _load_registry()
    now = time.time()
    to_remove = []

    for file_id, info in _registry.items():
        created = info.get("created_at", "")
        try:
            created_ts = datetime.fromisoformat(created).timestamp()
            if now - created_ts > max_age_hours * 3600:
                filepath = info.get("filepath", "")
                if os.path.exists(filepath):
                    os.remove(filepath)
                to_remove.append(file_id)
        except Exception:
            continue

    for fid in to_remove:
        del _registry[fid]

    if to_remove:
        _save_registry()
        logger.info(f"Cleaned up {len(to_remove)} old generated files")

    return len(to_remove)
